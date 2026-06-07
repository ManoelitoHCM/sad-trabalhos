"""
============================================================================
 SELEÇÃO DE BOLSISTAS COM ALGORITMO GENÉTICO - MICRODADOS DO ENEM
============================================================================
 Curso......: Sistemas de Informação
 Disciplina.: Sistemas de Apoio à Tomada de Decisão

 OBJETIVO
 --------
 Selecionar automaticamente o melhor grupo de 100 bolsistas que otimize,
 simultaneamente, três objetivos:
   1) Alta performance acadêmica  -> média das 5 notas        (peso 0.50)
   2) Diversidade socioeconômica  -> índice de diversidade    (peso 0.30)
      e racial
   3) Cobertura geográfica        -> nº de UFs representadas   (peso 0.20)

 MODELAGEM
 ---------
 - Cromossomo : um grupo de 100 candidatos (vetor de 100 índices distintos
                apontando para a base de dados filtrada).
 - Gene       : um candidato (uma linha da base).
 - Aptidão    : soma ponderada das três métricas, todas normalizadas em [0,1]:
                   f = 0.50*notas + 0.30*diversidade + 0.20*cobertura
 - Seleção    : torneio.
 - Crossover  : recombinação por união dos pais + amostragem de 100 distintos.
 - Mutação    : troca de uma fração dos genes por candidatos do pool.
 - Elitismo   : preserva os melhores indivíduos a cada geração.

 USO
 ---
   # 1) Com a amostra sintética (para teste rápido):
   python gerar_amostra_sintetica.py --n 5000 --saida amostra_enem.csv
   python algoritmo_genetico_bolsas.py --dados amostra_enem.csv

   # 2) ENEM 2024 — modo recomendado quando se exige 2024 (SÓ o RESULTADOS):
   python algoritmo_genetico_bolsas.py \
       --dados "microdados_enem_2024/DADOS/RESULTADOS_2024.csv" --amostra 100000

   # 3) ENEM <= 2023 (arquivo único e vinculado, com questionário completo):
   python algoritmo_genetico_bolsas.py --dados MICRODADOS_ENEM_2023.csv --amostra 100000

 OBSERVAÇÃO SOBRE O FORMATO 2024
 -------------------------------
 A partir de 2024 o INEP reestruturou os microdados (LGPD): as NOTAS ficam em
 RESULTADOS_2024.csv e o QUESTIONÁRIO socioeconômico (renda, cor/raça,
 escolaridade da mãe) em PARTICIPANTES_2024.csv, SEM chave de ligação entre eles.
 Como não dá para unir nota e perfil por participante, ao rodar com o RESULTADOS
 o script entra no MODO PROXY 2024: mede a diversidade pela dependência
 administrativa e pela localização da escola (proxies socioeconômicos),
 restrito aos participantes que declararam escola (concluintes).
 As notas (incl. NU_NOTA_REDACAO) e a cobertura por UF (SG_UF_PROVA) são detectadas
 automaticamente. O modo de dois arquivos (--pasta) é mantido apenas para
 hipotéticas bases vinculadas e avisa quando a junção é impossível.

 SAÍDAS (na pasta ./saida)
 -------
   - grupo_ideal.csv             -> os 100 bolsistas selecionados
   - convergencia.png            -> evolução da aptidão por geração
   - notas.png                   -> distribuição das notas do grupo
   - diversidade.png             -> composição socioeconômica e racial
   - distribuicao_geografica.png -> bolsistas por UF
   - Relatorio_Bolsas_AG.docx    -> relatório completo com resultados
============================================================================
"""

import argparse
import glob
import os
import random
import re
import time

import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")            # backend sem interface gráfica (salva PNGs)
import matplotlib.pyplot as plt


# ----------------------------------------------------------------------------
# CONFIGURAÇÃO (atributos relevantes definidos no enunciado)
# ----------------------------------------------------------------------------
COLUNAS_NOTAS = ['NU_NOTA_CN', 'NU_NOTA_CH', 'NU_NOTA_LC', 'NU_NOTA_MT', 'NU_NOTA_RED']
NOMES_AREAS   = {'NU_NOTA_CN': 'C. Natureza', 'NU_NOTA_CH': 'C. Humanas',
                 'NU_NOTA_LC': 'Linguagens', 'NU_NOTA_MT': 'Matemática',
                 'NU_NOTA_RED': 'Redação'}
COLUNAS_SOCIO = ['Q006', 'Q002', 'TP_ESCOLA', 'TP_COR_RACA']
NOMES_SOCIO   = {'Q006': 'Renda familiar', 'Q002': 'Escolaridade da mãe',
                 'TP_ESCOLA': 'Tipo de escola', 'TP_COR_RACA': 'Cor/raça'}
COLUNA_UF     = 'SG_UF_RESIDENCIA'

# ----------------------------------------------------------------------------
# DETECÇÃO DE COLUNAS (o formato 2024 mudou os arquivos e alguns nomes).
# Os nomes das NOTAS e dos códigos Q00x são estáveis entre as edições; a coluna
# de UF é a que mais varia, então testamos vários nomes em ordem de prioridade.
# Se necessário, acrescente nomes alternativos a estas listas.
# ----------------------------------------------------------------------------
CANDIDATOS_UF       = ['SG_UF_RESIDENCIA', 'SG_UF_PROVA', 'SG_UF_ESC',
                       'SG_UF_NASCIMENTO']
CANDIDATOS_PRESENCA = ['TP_PRESENCA_CN', 'TP_PRESENCA_CH',
                       'TP_PRESENCA_LC', 'TP_PRESENCA_MT']
CHAVE_JUNCAO        = 'NU_INSCRICAO'   # liga PARTICIPANTES e RESULTADOS

# Nomes-candidatos de cada nota (o 1º que existir é usado e renomeado para o
# nome canônico à esquerda). ATENÇÃO: nos microdados reais a redação é
# 'NU_NOTA_REDACAO' (o enunciado abrevia como 'NU_NOTA_RED').
NOTAS_CANDIDATOS = {
    'NU_NOTA_CN':  ['NU_NOTA_CN'],
    'NU_NOTA_CH':  ['NU_NOTA_CH'],
    'NU_NOTA_LC':  ['NU_NOTA_LC'],
    'NU_NOTA_MT':  ['NU_NOTA_MT'],
    'NU_NOTA_RED': ['NU_NOTA_RED', 'NU_NOTA_REDACAO'],
}

# Rótulos legíveis para os gráficos de diversidade
ROTULOS_ESCOLA = {1: 'Não respondeu', 2: 'Pública', 3: 'Privada'}
ROTULOS_RACA   = {0: 'Não declarado', 1: 'Branca', 2: 'Preta', 3: 'Parda',
                  4: 'Amarela', 5: 'Indígena', 6: 'Não dispõe'}
# Proxies disponíveis no RESULTADOS do ENEM 2024 (codificação padrão INEP).
ROTULOS_DEP_ADM = {1: 'Federal', 2: 'Estadual', 3: 'Municipal', 4: 'Privada'}
ROTULOS_LOC_ESC = {1: 'Urbana', 2: 'Rural'}

# --- Perfil de diversidade (definido em tempo de carga conforme as colunas) ---
# COLUNAS_SOCIO, os gráficos e o texto do relatório são ajustados de acordo.
DIVERSIDADE_PLOTS = [                       # (coluna, título, rótulos|None)
    ('TP_COR_RACA', 'Cor/Raça', ROTULOS_RACA),
    ('Q006', 'Renda familiar (Q006)', None),
    ('TP_ESCOLA', 'Tipo de escola', ROTULOS_ESCOLA),
    ('Q002', 'Escolaridade da mãe (Q002)', None),
]
DESC_DIVERSIDADE = ('socioeconômica e racial (renda, escolaridade da mãe, '
                    'tipo de escola e cor/raça)')
MODO_DIVERSIDADE = 'questionario'           # 'questionario' ou 'proxy_2024'

NUM_UF_BRASIL = 27               # 26 estados + Distrito Federal
NOTA_MAXIMA   = 1000.0           # teto teórico das notas do ENEM (normalização)

# Pesos da função de aptidão (conforme enunciado)
PESO_NOTAS        = 0.50
PESO_DIVERSIDADE  = 0.30
PESO_COBERTURA    = 0.20

# Parâmetros do algoritmo genético (conforme enunciado + operadores padrão)
TAM_GRUPO        = 100           # cromossomo = 100 bolsistas
TAM_POPULACAO    = 20            # 20 grupos na população inicial
NUM_GERACOES     = 100           # 100 gerações
TAM_TORNEIO      = 3             # nº de competidores na seleção por torneio
TAXA_MUTACAO     = 0.20          # prob. de um filho sofrer mutação
FRACAO_MUTACAO   = 0.05          # fração de genes trocados na mutação (5%)
NUM_ELITE        = 2             # melhores preservados a cada geração


# ============================================================================
# 1) PREPARAÇÃO DOS DADOS
# ============================================================================
def _detectar_formato(caminho):
    """Detecta (codificação, separador) lendo apenas o início do arquivo.
    O ENEM 2024 vem em UTF-8 (às vezes com BOM); edições antigas, em latin-1."""
    enc = 'utf-8-sig'
    try:
        with open(caminho, encoding='utf-8-sig') as f:
            primeira = f.readline()
            f.read(500_000)            # amostra para validar a codificação
    except UnicodeDecodeError:
        enc = 'latin-1'
        with open(caminho, encoding='latin-1') as f:
            primeira = f.readline()
    # separador mais frequente na linha de cabeçalho
    sep = max([';', ',', '\t', '|'], key=lambda s: primeira.count(s))
    return enc, sep


def _limpar_colunas(df):
    """Remove BOM e espaços dos nomes das colunas (robustez de leitura)."""
    df.columns = [str(c).replace('\ufeff', '').strip() for c in df.columns]
    return df


def _ler_csv(caminho, **kwargs):
    """Lê um CSV detectando codificação e separador, e limpando os nomes das
    colunas. Faz fallback para latin-1 caso a leitura UTF-8 falhe no meio."""
    enc, sep = _detectar_formato(caminho)
    try:
        df = pd.read_csv(caminho, sep=sep, encoding=enc, **kwargs)
    except UnicodeDecodeError:
        df = pd.read_csv(caminho, sep=sep, encoding='latin-1', **kwargs)
    return _limpar_colunas(df)


def _ler_cabecalho(caminho):
    """Lê apenas a 1ª linha (cabeçalho) do CSV — rápido e sem custo de memória."""
    return _ler_csv(caminho, nrows=0).columns.tolist()


def _localizar(cabecalho, candidatos, regex=None):
    """Acha no cabeçalho a 1ª coluna que casa com a lista (por prioridade) ou
    com o padrão regex. Retorna None se não encontrar."""
    for nome in candidatos:
        if nome in cabecalho:
            return nome
    if regex:
        for col in cabecalho:
            if re.search(regex, col):
                return col
    return None


def _mapear_notas(cabecalho):
    """Mapeia os nomes REAIS das notas no cabeçalho para os nomes CANÔNICOS.
    Ex.: {'NU_NOTA_REDACAO': 'NU_NOTA_RED', 'NU_NOTA_CN': 'NU_NOTA_CN', ...}."""
    mapa = {}
    for canonico, candidatos in NOTAS_CANDIDATOS.items():
        real = _localizar(cabecalho, candidatos)
        if real is not None:
            mapa[real] = canonico
    return mapa


def _definir_perfil_diversidade(cabecalho):
    """Escolhe os atributos do índice de diversidade conforme as colunas
    disponíveis no arquivo. Retorna (colunas, plots, descricao, modo).

    - 'questionario': arquivo com o questionário socioeconômico (sintético ou
      ENEM <= 2023) -> usa renda, escolaridade da mãe, tipo de escola e cor/raça.
    - 'proxy_2024'  : RESULTADOS do ENEM 2024 -> renda/raça/escolaridade da mãe
      NÃO estão neste arquivo (ficam no PARTICIPANTES, desvinculado), então a
      diversidade é medida por proxies da escola (dependência e localização).
    """
    tem = lambda c: c in cabecalho
    if all(tem(c) for c in ['Q006', 'Q002', 'TP_ESCOLA', 'TP_COR_RACA']):
        colunas = ['Q006', 'Q002', 'TP_ESCOLA', 'TP_COR_RACA']
        plots = [
            ('TP_COR_RACA', 'Cor/Raça', ROTULOS_RACA),
            ('Q006', 'Renda familiar (Q006)', None),
            ('TP_ESCOLA', 'Tipo de escola', ROTULOS_ESCOLA),
            ('Q002', 'Escolaridade da mãe (Q002)', None),
        ]
        desc = ('socioeconômica e racial (renda Q006, escolaridade da mãe Q002, '
                'tipo de escola e cor/raça)')
        return colunas, plots, desc, 'questionario'

    if tem('TP_DEPENDENCIA_ADM_ESC') or tem('TP_LOCALIZACAO_ESC'):
        colunas, plots = [], []
        if tem('TP_DEPENDENCIA_ADM_ESC'):
            colunas.append('TP_DEPENDENCIA_ADM_ESC')
            plots.append(('TP_DEPENDENCIA_ADM_ESC', 'Dependência da escola', ROTULOS_DEP_ADM))
        if tem('TP_LOCALIZACAO_ESC'):
            colunas.append('TP_LOCALIZACAO_ESC')
            plots.append(('TP_LOCALIZACAO_ESC', 'Localização da escola', ROTULOS_LOC_ESC))
        desc = ('por dependência administrativa (federal/estadual/municipal/'
                'privada) e localização (urbana/rural) da escola — PROXIES '
                'adotados porque renda, cor/raça e escolaridade da mãe não '
                'constam no RESULTADOS dos microdados 2024')
        return colunas, plots, desc, 'proxy_2024'

    return None


def _msg_coluna_ausente(rotulo, candidatos, cab_p, cab_r):
    def amostra_cols(cab):
        return ', '.join(cab[:30]) + (' ...' if len(cab) > 30 else '')
    return (f"Não encontrei a(s) coluna(s) de {rotulo}.\n"
            f"  Procurei por: {candidatos}\n"
            f"  Colunas disponíveis (PARTICIPANTES): {amostra_cols(cab_p)}\n"
            f"  Colunas disponíveis (RESULTADOS):    {amostra_cols(cab_r)}\n"
            f"  -> Ajuste as listas no bloco CONFIGURAÇÃO/DETECÇÃO do script.")


def _finalizar_pool(df, amostra, seed):
    """Validação final + amostragem, compartilhada pelos dois carregadores."""
    obrigatorias = COLUNAS_NOTAS + COLUNAS_SOCIO + [COLUNA_UF]
    faltando = [c for c in obrigatorias if c not in df.columns]
    if faltando:
        raise ValueError(f"Após o processamento, ainda faltam colunas: {faltando}")

    antes = len(df)
    df = df.dropna(subset=COLUNAS_NOTAS).reset_index(drop=True)
    if len(df) != antes:
        print(f"      Candidatos com todas as notas válidas: {len(df)} de {antes}")

    if amostra is not None and amostra < len(df):
        df = df.sample(n=amostra, random_state=seed).reset_index(drop=True)
        print(f"      Pool reduzido por amostragem para: {len(df)} candidatos")

    if len(df) < TAM_GRUPO:
        raise ValueError(f"Pool ({len(df)}) menor que o grupo exigido ({TAM_GRUPO}).")
    print(f"      Pool final: {len(df)} candidatos")
    return df


def carregar_dados(caminho, amostra=None, seed=42):
    """Carrega microdados em ARQUIVO ÚNICO (formato vinculado: amostra sintética,
    ENEM <= 2023, ou o RESULTADOS_2024.csv). Detecta automaticamente os nomes
    reais das notas (inclusive 'NU_NOTA_REDACAO'), a coluna de UF e o conjunto de
    atributos de diversidade disponível (questionário ou proxies de escola 2024)."""
    global COLUNAS_SOCIO, DIVERSIDADE_PLOTS, DESC_DIVERSIDADE, MODO_DIVERSIDADE

    print(f"[1/5] Lendo dados de '{caminho}' (arquivo único) ...")
    cab = _ler_cabecalho(caminho)

    # Notas (mapeia nomes reais -> canônicos; trata NU_NOTA_REDACAO).
    mapa_notas = _mapear_notas(cab)
    faltam_notas = [c for c in COLUNAS_NOTAS if c not in mapa_notas.values()]
    if faltam_notas:
        proc = [NOTAS_CANDIDATOS[c] for c in faltam_notas]
        raise ValueError(_msg_coluna_ausente(f"notas {faltam_notas}", proc, cab, cab))

    # UF (SG_UF_RESIDENCIA, SG_UF_PROVA, ...).
    uf_col = _localizar(cab, CANDIDATOS_UF, regex=r'^SG_UF')
    if uf_col is None:
        raise ValueError(_msg_coluna_ausente("UF (estado)", CANDIDATOS_UF, cab, cab))

    # Diversidade: escolhe o conjunto de atributos conforme as colunas presentes.
    perfil = _definir_perfil_diversidade(cab)
    if perfil is None:
        raise ValueError(_msg_coluna_ausente(
            "diversidade (Q006/Q002/TP_ESCOLA/TP_COR_RACA ou "
            "TP_DEPENDENCIA_ADM_ESC/TP_LOCALIZACAO_ESC)", ['Q006', '...'], cab, cab))
    COLUNAS_SOCIO, DIVERSIDADE_PLOTS, DESC_DIVERSIDADE, MODO_DIVERSIDADE = perfil
    if MODO_DIVERSIDADE == 'proxy_2024':
        print("      [info] Diversidade no MODO PROXY 2024: renda, cor/raça e "
              "escolaridade da mãe\n             não constam no RESULTADOS 2024; "
              "usando dependência e localização da escola.")
    print(f"      Atributos de diversidade: {COLUNAS_SOCIO}")

    colunas = list(mapa_notas.keys()) + COLUNAS_SOCIO + [uf_col]
    df = _ler_csv(caminho, usecols=lambda c: c in colunas, low_memory=False)
    df = df.rename(columns=mapa_notas)            # notas -> nomes canônicos
    for c in COLUNAS_NOTAS:                        # garante notas numéricas
        df[c] = pd.to_numeric(df[c], errors='coerce')
    if uf_col != COLUNA_UF:
        print(f"      [info] Coluna de UF detectada: '{uf_col}' "
              f"(o enunciado cita '{COLUNA_UF}'; será usada como o estado do candidato).")
        df = df.rename(columns={uf_col: COLUNA_UF})
    if 'NU_NOTA_REDACAO' in mapa_notas:
        print("      [info] Coluna de redação detectada: 'NU_NOTA_REDACAO'.")

    # No modo proxy 2024, os atributos de escola só existem para quem declarou
    # escola (concluintes). Mantemos apenas esses, para a diversidade ter sentido.
    if MODO_DIVERSIDADE == 'proxy_2024':
        for c in COLUNAS_SOCIO:
            df[c] = pd.to_numeric(df[c], errors='coerce')
        antes = len(df)
        df = df.dropna(subset=COLUNAS_SOCIO)
        for c in COLUNAS_SOCIO:
            df[c] = df[c].astype(int)
        print(f"      Participantes com escola declarada (proxies válidos): "
              f"{len(df)} de {antes}")

    return _finalizar_pool(df, amostra, seed)


def carregar_microdados_2024(caminho_participantes, caminho_resultados,
                             amostra=None, seed=42):
    """Carrega o formato ENEM 2024 (DOIS arquivos) e os une por NU_INSCRICAO.

    A partir de 2024 o INEP reestruturou os microdados: as notas ficam em
    RESULTADOS_2024.csv e os dados socioeconômicos/demográficos em
    PARTICIPANTES_2024.csv. Esta função detecta em qual arquivo está cada
    coluna, lê apenas o necessário (para caber na memória) e junta os dois.
    """
    print("[1/5] Lendo microdados no formato 2024 (dois arquivos) ...")
    cab_p = _ler_cabecalho(caminho_participantes)
    cab_r = _ler_cabecalho(caminho_resultados)

    if CHAVE_JUNCAO not in cab_p or CHAVE_JUNCAO not in cab_r:
        raise ValueError(
            f"Não há uma coluna de junção '{CHAVE_JUNCAO}' comum aos dois arquivos.\n"
            f"  Colunas em PARTICIPANTES: {cab_p}\n"
            f"  Colunas em RESULTADOS:    {cab_r}\n"
            f"  ATENÇÃO: a partir de 2024 o INEP DESVINCULA PARTICIPANTES e RESULTADOS\n"
            f"  (chaves diferentes, ex.: NU_INSCRICAO x NU_SEQUENCIAL) por anonimização\n"
            f"  (LGPD). Isso IMPEDE unir nota e perfil socioeconômico por participante.\n"
            f"  -> Opção A (recomendada): edição com arquivo ÚNICO e vinculado (ENEM 2023):\n"
            f"       python algoritmo_genetico_bolsas.py --dados MICRODADOS_ENEM_2023.csv --amostra 100000\n"
            f"  -> Opção B (manter 2024): rode SÓ no RESULTADOS, com diversidade por\n"
            f"     proxies de escola (dependência e localização):\n"
            f"       python algoritmo_genetico_bolsas.py --dados RESULTADOS_2024.csv --amostra 100000")

    def onde(coluna):
        if coluna in cab_r:
            return 'R'
        if coluna in cab_p:
            return 'P'
        return None

    # --- Notas (geralmente em RESULTADOS) ---
    faltam_notas = [c for c in COLUNAS_NOTAS if onde(c) is None]
    if faltam_notas:
        raise ValueError(_msg_coluna_ausente("notas", faltam_notas, cab_p, cab_r))

    # --- Presença (geralmente em RESULTADOS); opcional ---
    presenca = [c for c in CANDIDATOS_PRESENCA if onde(c) is not None]

    # --- Socioeconômicos (geralmente em PARTICIPANTES) ---
    socio_em = {}
    for c in COLUNAS_SOCIO:
        loc = onde(c)
        if loc is None:
            raise ValueError(_msg_coluna_ausente(f"socioeconômico '{c}'",
                                                 [c], cab_p, cab_r))
        socio_em[c] = loc

    # --- UF (pode estar em qualquer um dos arquivos) ---
    uf_p = _localizar(cab_p, CANDIDATOS_UF, regex=r'^SG_UF')
    uf_r = _localizar(cab_r, CANDIDATOS_UF, regex=r'^SG_UF')
    uf_col = uf_p or uf_r
    if uf_col is None:
        raise ValueError(_msg_coluna_ausente("UF (estado)", CANDIDATOS_UF, cab_p, cab_r))
    uf_arq = 'P' if uf_p else 'R'
    if uf_col != COLUNA_UF:
        print(f"      [info] Coluna de UF detectada: '{uf_col}' (arquivo "
              f"{'PARTICIPANTES' if uf_arq == 'P' else 'RESULTADOS'}). O enunciado "
              f"cita '{COLUNA_UF}'; '{uf_col}' será usada como o estado do candidato.")

    # --- Define quais colunas ler de cada arquivo (roteia pela detecção) ---
    cols_R, cols_P = [CHAVE_JUNCAO], [CHAVE_JUNCAO]
    def rotear(coluna, arquivo):
        (cols_R if arquivo == 'R' else cols_P).append(coluna)
    for c in COLUNAS_NOTAS:                       # notas (onde estiverem)
        rotear(c, onde(c))
    for c in presenca:                            # presença (onde estiver)
        rotear(c, onde(c))
    for c, loc in socio_em.items():               # socioeconômicos
        rotear(c, loc)
    rotear(uf_col, uf_arq)                         # UF
    cols_R = list(dict.fromkeys(cols_R))          # remove duplicatas mantendo ordem
    cols_P = list(dict.fromkeys(cols_P))

    # --- Lê RESULTADOS (arquivo grande): só colunas necessárias e já filtra ---
    print(f"      RESULTADOS -> lendo colunas {cols_R}")
    res = _ler_csv(caminho_resultados, usecols=lambda c: c in cols_R, low_memory=False)
    notas_em_res = [c for c in COLUNAS_NOTAS if c in res.columns]
    for c in notas_em_res:                                  # notas -> numérico
        res[c] = pd.to_numeric(res[c], errors='coerce')
    for c in [p for p in presenca if p in res.columns]:     # filtra presentes
        res = res[res[c] == 1]
    if notas_em_res:                                        # filtro antecipado
        res = res.dropna(subset=notas_em_res)
    print(f"      RESULTADOS após filtro de presença/notas: {len(res)} linhas")

    # --- Lê PARTICIPANTES: só colunas necessárias ---
    print(f"      PARTICIPANTES -> lendo colunas {cols_P}")
    par = _ler_csv(caminho_participantes, usecols=lambda c: c in cols_P, low_memory=False)
    for c in [n for n in COLUNAS_NOTAS if n in par.columns]:  # caso raro: nota aqui
        par[c] = pd.to_numeric(par[c], errors='coerce')

    # --- JOIN por NU_INSCRICAO (inner: só quem tem nota E cadastro) ---
    df = res.merge(par, on=CHAVE_JUNCAO, how='inner')
    print(f"      Após o JOIN por {CHAVE_JUNCAO}: {len(df)} candidatos")

    if uf_col != COLUNA_UF:                        # padroniza o nome da UF
        df = df.rename(columns={uf_col: COLUNA_UF})

    return _finalizar_pool(df, amostra, seed)


def _resolver_arquivos_pasta(pasta):
    """Dada a pasta DADOS, localiza os arquivos PARTICIPANTES e RESULTADOS."""
    def achar(padrao):
        encontrados = glob.glob(os.path.join(pasta, padrao))
        return encontrados[0] if encontrados else None
    part = achar('PARTICIPANTES*.csv') or achar('*PARTICIPANTES*.csv')
    resu = achar('RESULTADOS*.csv') or achar('*RESULTADOS*.csv')
    if not part or not resu:
        raise ValueError(
            f"Não localizei PARTICIPANTES*.csv e/ou RESULTADOS*.csv em '{pasta}'.\n"
            f"  Arquivos .csv encontrados: {glob.glob(os.path.join(pasta, '*.csv'))}")
    return part, resu


def preparar_estruturas(df):
    """Pré-computa arrays NumPy para a função de aptidão ser rápida no laço.

    Retorna um dicionário com:
      - notas        : matriz (N x 5) de notas
      - socio_codes  : dict coluna -> (códigos inteiros, nº de categorias)
      - uf_codes     : códigos inteiros das UFs
    """
    notas = df[COLUNAS_NOTAS].to_numpy(dtype=float)

    socio_codes = {}
    for col in COLUNAS_SOCIO:
        codigos, categorias = pd.factorize(df[col])
        socio_codes[col] = (codigos.astype(np.int64), len(categorias))

    uf_codes, _ = pd.factorize(df[COLUNA_UF])
    return {'notas': notas, 'socio_codes': socio_codes,
            'uf_codes': uf_codes.astype(np.int64)}


# ============================================================================
# 2) FUNÇÃO DE APTIDÃO E SUAS COMPONENTES
# ============================================================================
def _entropia_normalizada(codigos, n_categorias):
    """Entropia de Shannon normalizada em [0,1].
    0 = todos iguais (sem diversidade); 1 = distribuição uniforme (máxima)."""
    if n_categorias <= 1:
        return 0.0
    contagem = np.bincount(codigos, minlength=n_categorias)
    contagem = contagem[contagem > 0]
    p = contagem / contagem.sum()
    h = -np.sum(p * np.log(p))
    return float(h / np.log(n_categorias))


def componentes_aptidao(indices, est):
    """Calcula as três métricas (todas em [0,1]) para um grupo de candidatos."""
    # (1) Performance acadêmica: média das 5 notas, normalizada por 1000.
    media_notas = est['notas'][indices].mean() / NOTA_MAXIMA

    # (2) Diversidade: média das entropias normalizadas dos 4 atributos socio.
    entropias = [
        _entropia_normalizada(codigos[indices], n_cat)
        for (codigos, n_cat) in est['socio_codes'].values()
    ]
    diversidade = float(np.mean(entropias))

    # (3) Cobertura geográfica: fração das 27 UFs representadas no grupo.
    cobertura = len(np.unique(est['uf_codes'][indices])) / NUM_UF_BRASIL

    return media_notas, diversidade, cobertura


def aptidao(indices, est):
    """Aptidão final: soma ponderada das três componentes."""
    notas, diversidade, cobertura = componentes_aptidao(indices, est)
    valor = (PESO_NOTAS * notas
             + PESO_DIVERSIDADE * diversidade
             + PESO_COBERTURA * cobertura)
    return valor, (notas, diversidade, cobertura)


# ============================================================================
# 3) OPERADORES GENÉTICOS
# ============================================================================
def populacao_inicial(n_pool, rng):
    """Cria TAM_POPULACAO grupos, cada um com 100 candidatos distintos."""
    return [rng.choice(n_pool, size=TAM_GRUPO, replace=False)
            for _ in range(TAM_POPULACAO)]


def selecao_torneio(populacao, aptidoes, rng):
    """Seleciona um indivíduo: sorteia TAM_TORNEIO competidores e vence o de
    maior aptidão."""
    competidores = rng.choice(len(populacao), size=TAM_TORNEIO, replace=False)
    melhor = max(competidores, key=lambda i: aptidoes[i])
    return populacao[melhor]


def crossover(pai_a, pai_b, rng):
    """Recombina dois pais: une os candidatos dos dois e sorteia 100 distintos.
    Garante um filho válido (sem repetição e com exatamente 100 genes)."""
    uniao = np.union1d(pai_a, pai_b)              # entre 100 e 200 candidatos
    return rng.choice(uniao, size=TAM_GRUPO, replace=False)


def mutacao(cromossomo, n_pool, rng):
    """Com probabilidade TAXA_MUTACAO, troca uma fração dos genes por novos
    candidatos do pool (mantendo o grupo sem repetições)."""
    if rng.random() >= TAXA_MUTACAO:
        return cromossomo
    filho = cromossomo.copy()
    presentes = set(filho.tolist())
    n_trocas = max(1, int(FRACAO_MUTACAO * TAM_GRUPO))
    posicoes = rng.choice(TAM_GRUPO, size=n_trocas, replace=False)
    for pos in posicoes:
        novo = int(rng.integers(n_pool))
        while novo in presentes:                  # evita duplicar candidato
            novo = int(rng.integers(n_pool))
        presentes.discard(int(filho[pos]))
        presentes.add(novo)
        filho[pos] = novo
    return filho


# ============================================================================
# 4) LAÇO PRINCIPAL DO ALGORITMO GENÉTICO
# ============================================================================
def algoritmo_genetico(est, seed=42, verbose=True):
    rng = np.random.default_rng(seed)
    n_pool = len(est['uf_codes'])

    populacao = populacao_inicial(n_pool, rng)
    historico = {'melhor': [], 'media': [], 'notas': [],
                 'diversidade': [], 'cobertura': []}

    melhor_global, melhor_apt = None, -np.inf
    if verbose:
        print(f"[2/5] Executando AG: {TAM_POPULACAO} grupos x {NUM_GERACOES} gerações ...")
        print("      Ger | melhor aptidão | notas  diversid.  cobertura")

    for ger in range(NUM_GERACOES):
        # --- Avaliação ---
        resultados = [aptidao(ind, est) for ind in populacao]
        aptidoes = np.array([r[0] for r in resultados])

        ordem = np.argsort(aptidoes)[::-1]            # do melhor para o pior
        if aptidoes[ordem[0]] > melhor_apt:
            melhor_apt = aptidoes[ordem[0]]
            melhor_global = populacao[ordem[0]].copy()

        _, comp = resultados[ordem[0]]
        historico['melhor'].append(aptidoes[ordem[0]])
        historico['media'].append(aptidoes.mean())
        historico['notas'].append(comp[0])
        historico['diversidade'].append(comp[1])
        historico['cobertura'].append(comp[2])

        if verbose and (ger % 10 == 0 or ger == NUM_GERACOES - 1):
            print(f"      {ger:3d} | {aptidoes[ordem[0]]:.6f}     | "
                  f"{comp[0]:.4f}  {comp[1]:.4f}    {comp[2]:.4f}")

        # --- Nova geração ---
        nova = [populacao[ordem[i]].copy() for i in range(NUM_ELITE)]  # elitismo
        while len(nova) < TAM_POPULACAO:
            pai_a = selecao_torneio(populacao, aptidoes, rng)
            pai_b = selecao_torneio(populacao, aptidoes, rng)
            filho = crossover(pai_a, pai_b, rng)
            filho = mutacao(filho, n_pool, rng)
            nova.append(filho)
        populacao = nova

    if verbose:
        print(f"      Melhor aptidão final: {melhor_apt:.6f}")
    return melhor_global, melhor_apt, historico


# ============================================================================
# 5) ANÁLISE DOS RESULTADOS (GRÁFICOS + EXPORTAÇÃO)
# ============================================================================
def grafico_convergencia(historico, caminho):
    fig, ax = plt.subplots(figsize=(9, 5))
    ax.plot(historico['melhor'], label='Melhor aptidão', linewidth=2.2)
    ax.plot(historico['media'], label='Aptidão média', linewidth=1.5, alpha=0.8)
    ax.plot(historico['notas'], '--', label='Componente notas (0.5)', alpha=0.7)
    ax.plot(historico['diversidade'], '--', label='Componente diversidade (0.3)', alpha=0.7)
    ax.plot(historico['cobertura'], '--', label='Componente cobertura (0.2)', alpha=0.7)
    ax.set_title('Convergência do Algoritmo Genético')
    ax.set_xlabel('Geração'); ax.set_ylabel('Valor (normalizado)')
    ax.legend(loc='lower right', fontsize=8); ax.grid(alpha=0.3)
    fig.tight_layout(); fig.savefig(caminho, dpi=130); plt.close(fig)


def grafico_notas(grupo, caminho):
    dados = [grupo[c].values for c in COLUNAS_NOTAS]
    rotulos = [NOMES_AREAS[c] for c in COLUNAS_NOTAS]
    fig, ax = plt.subplots(figsize=(9, 5))
    bp = ax.boxplot(dados, tick_labels=rotulos, patch_artist=True, showmeans=True)
    for caixa in bp['boxes']:
        caixa.set(facecolor='#4C72B0', alpha=0.6)
    ax.set_title('Distribuição das Notas do Grupo Selecionado')
    ax.set_ylabel('Nota'); ax.grid(axis='y', alpha=0.3)
    fig.tight_layout(); fig.savefig(caminho, dpi=130); plt.close(fig)


def grafico_diversidade(grupo, caminho):
    plots = DIVERSIDADE_PLOTS
    n = len(plots)
    ncols = 2 if n > 1 else 1
    nrows = (n + ncols - 1) // ncols
    fig, axs = plt.subplots(nrows, ncols, figsize=(5.5 * ncols, 4 * nrows),
                            squeeze=False)
    eixos = axs.flatten()

    def barras(ax, serie, titulo, rotulos=None):
        cont = serie.value_counts().sort_index()
        nomes = [str(rotulos.get(i, i)) if rotulos else str(i) for i in cont.index]
        ax.bar(nomes, cont.values, color='#55A868', alpha=0.85)
        ax.set_title(titulo, fontsize=10)
        ax.tick_params(axis='x', labelrotation=30, labelsize=8)
        for i, v in enumerate(cont.values):
            ax.text(i, v, str(v), ha='center', va='bottom', fontsize=7)

    for ax, (coluna, titulo, rotulos) in zip(eixos, plots):
        barras(ax, grupo[coluna], titulo, rotulos)
    for ax in eixos[n:]:                          # esconde subplots não usados
        ax.axis('off')

    sufixo = ('Socioeconômica e Racial' if MODO_DIVERSIDADE == 'questionario'
              else 'do Grupo (proxies de escola — ENEM 2024)')
    fig.suptitle(f'Diversidade {sufixo}', fontsize=13)
    fig.tight_layout(rect=[0, 0, 1, 0.96]); fig.savefig(caminho, dpi=130); plt.close(fig)


def grafico_geografia(grupo, caminho):
    cont = grupo[COLUNA_UF].value_counts().sort_values(ascending=False)
    fig, ax = plt.subplots(figsize=(11, 5))
    ax.bar(cont.index.astype(str), cont.values, color='#C44E52', alpha=0.85)
    ax.set_title(f'Distribuição Geográfica dos Bolsistas '
                 f'({cont.shape[0]} de {NUM_UF_BRASIL} UFs)')
    ax.set_xlabel('UF'); ax.set_ylabel('Nº de bolsistas')
    ax.grid(axis='y', alpha=0.3)
    fig.tight_layout(); fig.savefig(caminho, dpi=130); plt.close(fig)


def gerar_graficos(grupo, historico, pasta):
    caminhos = {
        'convergencia': os.path.join(pasta, 'convergencia.png'),
        'notas':        os.path.join(pasta, 'notas.png'),
        'diversidade':  os.path.join(pasta, 'diversidade.png'),
        'geografia':    os.path.join(pasta, 'distribuicao_geografica.png'),
    }
    print("[4/5] Gerando gráficos ...")
    grafico_convergencia(historico, caminhos['convergencia'])
    grafico_notas(grupo, caminhos['notas'])
    grafico_diversidade(grupo, caminhos['diversidade'])
    grafico_geografia(grupo, caminhos['geografia'])
    return caminhos


def resumo_resultados(grupo, apt, comp):
    """Monta um dicionário de métricas-resumo do grupo escolhido."""
    notas, diversidade, cobertura = comp
    medias_area = {NOMES_AREAS[c]: round(grupo[c].mean(), 1) for c in COLUNAS_NOTAS}
    return {
        'aptidao': round(apt, 6),
        'media_geral': round(grupo[COLUNAS_NOTAS].mean().mean(), 1),
        'medias_area': medias_area,
        'comp_notas': round(notas, 4),
        'comp_diversidade': round(diversidade, 4),
        'comp_cobertura': round(cobertura, 4),
        'ufs_cobertas': int(grupo[COLUNA_UF].nunique()),
    }


# ============================================================================
# 6) GERAÇÃO DO RELATÓRIO (.docx)
# ============================================================================
def gerar_relatorio_docx(grupo, resumo, graficos, caminho, fonte_dados):
    """Cria o relatório completo em Word com metodologia + resultados reais."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    print("[5/5] Gerando relatório DOCX ...")
    doc = Document()
    # Garante que o elemento <w:zoom> tenha 'percent' (exigido pelo schema OOXML).
    _zoom = doc.settings.element.find(qn('w:zoom'))
    if _zoom is not None:
        _zoom.set(qn('w:percent'), '100')
    estilo = doc.styles['Normal']
    estilo.font.name = 'Calibri'; estilo.font.size = Pt(11)

    def titulo(txt, nivel=1):
        h = doc.add_heading(txt, level=nivel)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x3B, 0x73)
        return h

    def par(txt):
        return doc.add_paragraph(txt)

    # ---- Capa ----
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run('Seleção de Bolsistas com Algoritmo Genético')
    r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(0x1F, 0x3B, 0x73)
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run('Microdados do ENEM — Sistemas de Apoio à Tomada de Decisão')
    sr.italic = True; sr.font.size = Pt(12)
    doc.add_paragraph()

    # ---- 1. Introdução ----
    titulo('1. Introdução', 1)
    par('Uma fundação educacional privada deseja conceder bolsas de estudo a '
        'estudantes de alto potencial que realizaram o ENEM. O objetivo é '
        'maximizar o impacto social e acadêmico das bolsas, selecionando um '
        'grupo de 100 estudantes a partir de múltiplos critérios. Este relatório '
        'descreve a solução desenvolvida com um Algoritmo Genético (AG) que '
        'otimiza, de forma simultânea, três objetivos: desempenho acadêmico, '
        'diversidade socioeconômica e racial e cobertura geográfica nacional.')

    # ---- 2. Base de dados ----
    titulo('2. Base de Dados e Preparação', 1)
    par(f'Fonte utilizada nesta execução: {fonte_dados}. Os microdados oficiais '
        'do ENEM estão disponíveis no portal do INEP e utilizam separador ";" e '
        'codificação latin-1. Foram mantidas apenas as colunas relevantes para o '
        'problema e somente os candidatos com todas as notas válidas (participantes '
        'ausentes, cujas notas ficam em branco, são descartados).')
    par('Atributos utilizados:')
    for txt in ['Notas: NU_NOTA_CN, NU_NOTA_CH, NU_NOTA_LC, NU_NOTA_MT, NU_NOTA_RED.',
                'Socioeconômicos: Q006 (renda familiar), Q002 (escolaridade da mãe), '
                'TP_ESCOLA (tipo de escola) e TP_COR_RACA (cor/raça).',
                'Localidade: SG_UF_RESIDENCIA (unidade da federação de residência).']:
        doc.add_paragraph(txt, style='List Bullet')

    # ---- 3. Modelagem do AG ----
    titulo('3. Configuração do Algoritmo Genético', 1)
    titulo('3.1 Representação', 2)
    par('Cada cromossomo representa um grupo de 100 candidatos, codificado como '
        'um vetor de 100 índices distintos que apontam para linhas da base '
        'filtrada. Cada gene corresponde, portanto, a um candidato.')

    titulo('3.2 Função de Aptidão', 2)
    par('A aptidão é a soma ponderada de três componentes, todas normalizadas '
        'no intervalo [0, 1], conforme os pesos definidos no enunciado:')
    par('f = 0,50 × (média das notas) + 0,30 × (índice de diversidade) '
        '+ 0,20 × (cobertura regional)')
    for txt in ['Média das notas (peso 0,50): média das cinco notas do grupo, '
                'dividida por 1000 para normalizar.',
                f'Índice de diversidade (peso 0,30): média da entropia de Shannon '
                f'normalizada dos atributos de diversidade — {DESC_DIVERSIDADE}. '
                f'Entropia 0 indica ausência de diversidade; 1 indica distribuição '
                f'uniforme entre as categorias.',
                'Cobertura regional (peso 0,20): fração das 27 unidades da federação '
                '(26 estados + DF) representadas no grupo.']:
        doc.add_paragraph(txt, style='List Bullet')

    if MODO_DIVERSIDADE == 'proxy_2024':
        obs = doc.add_paragraph()
        r = obs.add_run('Observação metodológica (microdados 2024): ')
        r.bold = True
        obs.add_run(
            'a partir de 2024 o INEP passou a divulgar os microdados em arquivos '
            'separados (PARTICIPANTES e RESULTADOS) sem chave de ligação entre eles, '
            'por anonimização (LGPD). Como as notas estão no RESULTADOS e o '
            'questionário socioeconômico (renda, escolaridade da mãe) e a cor/raça '
            'estão no PARTICIPANTES, não é possível uni-los por participante. Por '
            'isso, o índice de diversidade foi medido pelos atributos disponíveis no '
            'RESULTADOS — a dependência administrativa (federal/estadual/municipal/'
            'privada) e a localização (urbana/rural) da escola —, que funcionam como '
            'proxies socioeconômicos. A análise restringe-se aos participantes que '
            'declararam escola (concluintes), para os quais esses atributos existem.')

    titulo('3.3 Operadores e Parâmetros', 2)
    par('Foram empregados operadores genéticos padrão:')
    for txt in ['Seleção por torneio (3 competidores por disputa).',
                'Crossover por união dos pais seguido de amostragem de 100 '
                'candidatos distintos, garantindo filhos válidos.',
                'Mutação que troca 5% dos genes por novos candidatos do conjunto, '
                'com probabilidade de 20% por indivíduo.',
                'Elitismo preservando os 2 melhores grupos a cada geração.']:
        doc.add_paragraph(txt, style='List Bullet')

    # Tabela de parâmetros
    par('Parâmetros de execução:')
    tab = doc.add_table(rows=1, cols=2); tab.style = 'Light Grid Accent 1'
    tab.rows[0].cells[0].text = 'Parâmetro'
    tab.rows[0].cells[1].text = 'Valor'
    parametros = [
        ('Tamanho do grupo (cromossomo)', str(TAM_GRUPO)),
        ('População inicial', str(TAM_POPULACAO)),
        ('Número de gerações', str(NUM_GERACOES)),
        ('Tamanho do torneio', str(TAM_TORNEIO)),
        ('Taxa de mutação', f'{TAXA_MUTACAO:.0%}'),
        ('Fração de genes mutados', f'{FRACAO_MUTACAO:.0%}'),
        ('Elitismo', str(NUM_ELITE)),
    ]
    for nome, valor in parametros:
        linha = tab.add_row().cells
        linha[0].text = nome; linha[1].text = valor

    # ---- 4. Resultados ----
    titulo('4. Análise dos Resultados', 1)
    par(f'A aptidão final do melhor grupo foi {resumo["aptidao"]}. As médias por '
        f'área de conhecimento e o desempenho geral do grupo escolhido são '
        f'apresentados a seguir.')
    tab2 = doc.add_table(rows=1, cols=2); tab2.style = 'Light Grid Accent 1'
    tab2.rows[0].cells[0].text = 'Métrica'
    tab2.rows[0].cells[1].text = 'Valor'
    metricas = [
        ('Aptidão final', str(resumo['aptidao'])),
        ('Média geral das notas', str(resumo['media_geral'])),
        ('Componente notas (norm.)', str(resumo['comp_notas'])),
        ('Componente diversidade (norm.)', str(resumo['comp_diversidade'])),
        ('Componente cobertura (norm.)', str(resumo['comp_cobertura'])),
        ('UFs representadas', f'{resumo["ufs_cobertas"]} de {NUM_UF_BRASIL}'),
    ]
    for c in COLUNAS_NOTAS:
        metricas.append((f'Média — {NOMES_AREAS[c]}', str(resumo['medias_area'][NOMES_AREAS[c]])))
    for nome, valor in metricas:
        linha = tab2.add_row().cells
        linha[0].text = nome; linha[1].text = valor

    def add_figura(caminho, legenda):
        doc.add_paragraph()
        doc.add_picture(caminho, width=Inches(6.0))
        leg = doc.add_paragraph(legenda); leg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in leg.runs:
            run.italic = True; run.font.size = Pt(9)

    titulo('4.1 Convergência', 2)
    par('A curva abaixo mostra a evolução da melhor aptidão e da aptidão média '
        'ao longo das gerações, além das três componentes da função objetivo.')
    add_figura(graficos['convergencia'], 'Figura 1 — Convergência do algoritmo genético.')

    titulo('4.2 Desempenho Acadêmico', 2)
    add_figura(graficos['notas'], 'Figura 2 — Distribuição das notas do grupo selecionado.')

    _sub43 = ('Diversidade Socioeconômica e Racial'
              if MODO_DIVERSIDADE == 'questionario'
              else 'Diversidade (proxies de escola — ENEM 2024)')
    _leg3 = ('Figura 3 — Composição socioeconômica e racial.'
             if MODO_DIVERSIDADE == 'questionario'
             else 'Figura 3 — Composição por dependência e localização da escola.')
    titulo(f'4.3 {_sub43}', 2)
    add_figura(graficos['diversidade'], _leg3)

    titulo('4.4 Distribuição Geográfica', 2)
    add_figura(graficos['geografia'], 'Figura 4 — Distribuição dos bolsistas por UF.')

    # ---- 5. Conclusão ----
    titulo('5. Conclusão', 1)
    par('O algoritmo genético mostrou-se adequado para um problema de seleção '
        'combinatória com múltiplos objetivos conflitantes. A formulação por soma '
        'ponderada permitiu equilibrar mérito acadêmico, diversidade e abrangência '
        'geográfica em uma única medida de aptidão, e o elitismo garantiu melhora '
        'monotônica da melhor solução ao longo das gerações. Os pesos (0,5 / 0,3 / '
        '0,2) podem ser ajustados pela fundação para refletir suas prioridades, e '
        'a abordagem pode ser estendida para uma otimização multiobjetivo (p. ex., '
        'NSGA-II) caso se deseje analisar o trade-off entre os objetivos por meio '
        'de uma fronteira de Pareto.')

    doc.save(caminho)
    return caminho


# ============================================================================
# MAIN
# ============================================================================
def main():
    ap = argparse.ArgumentParser(
        description="Algoritmo genético para seleção de bolsistas (ENEM).")
    # Modo 1 (arquivo único / amostra sintética):
    ap.add_argument('--dados', help="CSV único (amostra sintética ou formato antigo).")
    # Modo 2 (formato ENEM 2024 - dois arquivos):
    ap.add_argument('--pasta', help="Pasta DADOS do ENEM 2024 (acha PARTICIPANTES e "
                                    "RESULTADOS automaticamente).")
    ap.add_argument('--participantes', help="Caminho de PARTICIPANTES_2024.csv.")
    ap.add_argument('--resultados', help="Caminho de RESULTADOS_2024.csv.")
    ap.add_argument('--amostra', type=int, default=None,
                    help="Subamostra o pool para N candidatos (desempenho).")
    ap.add_argument('--saida', default='saida', help="Pasta de saída.")
    ap.add_argument('--seed', type=int, default=42, help="Semente aleatória.")
    args = ap.parse_args()

    random.seed(args.seed); np.random.seed(args.seed)
    os.makedirs(args.saida, exist_ok=True)
    t0 = time.time()

    # 1) Dados - escolhe o modo de carga conforme os argumentos informados.
    if args.pasta:
        part, resu = _resolver_arquivos_pasta(args.pasta)
        fonte = "Microdados ENEM 2024 (PARTICIPANTES + RESULTADOS)"
        df = carregar_microdados_2024(part, resu, amostra=args.amostra, seed=args.seed)
    elif args.participantes and args.resultados:
        fonte = "Microdados ENEM 2024 (PARTICIPANTES + RESULTADOS)"
        df = carregar_microdados_2024(args.participantes, args.resultados,
                                      amostra=args.amostra, seed=args.seed)
    elif args.dados:
        fonte = os.path.basename(args.dados)
        df = carregar_dados(args.dados, amostra=args.amostra, seed=args.seed)
    else:
        ap.error("Informe --pasta (ENEM 2024), OU --participantes e --resultados, "
                 "OU --dados (arquivo único).")

    est = preparar_estruturas(df)

    # 2-3) Algoritmo genético
    melhor, apt, historico = algoritmo_genetico(est, seed=args.seed)
    _, comp = aptidao(melhor, est)
    grupo = df.iloc[melhor].reset_index(drop=True)

    # Exporta grupo ideal
    caminho_csv = os.path.join(args.saida, 'grupo_ideal.csv')
    grupo.to_csv(caminho_csv, sep=';', encoding='utf-8-sig', index=False)
    print(f"[3/5] Grupo ideal salvo em '{caminho_csv}'")

    # 4) Gráficos + resumo
    graficos = gerar_graficos(grupo, historico, args.saida)
    resumo = resumo_resultados(grupo, apt, comp)

    # 5) Relatório
    caminho_doc = os.path.join(args.saida, 'Relatorio_Bolsas_AG.docx')
    gerar_relatorio_docx(grupo, resumo, graficos, caminho_doc, fonte_dados=fonte)

    print("\n================ RESUMO ================")
    print(f" Aptidão final........: {resumo['aptidao']}")
    print(f" Média geral das notas: {resumo['media_geral']}")
    print(f" Diversidade (norm.)..: {resumo['comp_diversidade']}")
    print(f" Cobertura............: {resumo['ufs_cobertas']}/{NUM_UF_BRASIL} UFs")
    print(f" Tempo total..........: {time.time() - t0:.1f}s")
    print("========================================")


if __name__ == '__main__':
    main()